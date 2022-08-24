# !/usr/bin/env python
# !-*- coding:utf-8 -*-
# !@Time   : 2022/5/23 14:52
# !@Author : DongHan Yang
# !@File   : 2.py
import docx
from openpyxl import Workbook
import os
import configparser


def write(file_pp, listst, choose):
    document = docx.Document(file_pp)
    tables = document.tables
    if choose == "*":
        table_indexs = [i for i in range(len(tables))]
    else:
        table_indexs = list(map(int, choose.split(',')))
    nums0, nums1 = 0, 0
    for index in table_indexs:
        lst = []
        table = tables[index]
        try:
            for i, j in listst:  # 处理word对应的单元格
                tx = table.cell(i, j).text
                if cf.get("config", "repalceSpace") == 'True':
                    tx = tx.replace(" ", "")
                lst.append(tx)
            worksheet.append(lst)
            nums0 += 1
        except Exception as e:
            print("-" * 50)
            print(f"第{nums0 + nums1}个表出错，请核对表内容")
            print(f"表{nums0 + nums1}错误信息：{e}")
            nums1 += 1
    return nums0, nums1


def read(file_pp, ceilTxt):
    document = docx.Document(file_pp)
    tables = document.tables
    # print(len(tables))
    ceilList = ceilTxt.split(",")
    table = tables[0]
    lst = ceilTxt.split(",")
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if cell.text in ceilList:
                # print(cell.text)
                i = ceilList.index(cell.text)
                if lst[i] in ceilList:
                    lst[i] = [row_index, col_index]
    return lst


def file_name(file_dir, gs):
    L, LL = [], []
    for root, dirs, files in os.walk(file_dir):
        for file in files:
            names = os.path.splitext(file)
            if names[1].split('.')[1] == gs:  # 想要保存的文件格式
                L.append(file)
                LL.append(root + '/' + file)
    print(f"符合要求进行提取的文件共{len(L)}个,为：")
    print(L)
    return LL


if __name__ == '__main__':
    print("请修改config.ini和阅读readme.txt; 如遇问题可邮件联系1198391037@qq.com")
    try:
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "Sheet1"

        # 创建配置类对象
        cf = configparser.ConfigParser()
        # 读取配置文件
        cf.read("config.ini", encoding="utf-8-sig")
        options = cf.options("config")
        path = cf.get("config", "path").replace("\\", "/")
        save_file = path + "/1.xlsx"
        types = cf.get("config", "doc")
        tableNums = cf.get("config", "tableNums")
        tableIndexs = cf.get("config", "tableIndexs")
        if tableIndexs != "True":
            tableIndexss = []
            ceilIndexs = cf.get("ceilIndex", "setIndexs")
            for t in ceilIndexs.split(";"):
                tableIndexss.append(list(map(int, t.split(','))))
        else:
            Demopath = cf.get("ceilIndex", "Demopath")
            tableTxt = cf.get("ceilIndex", "tableTxt")
            tableIndexss = read(Demopath, tableTxt)
        print(f"word定位的单元格位置为{tableIndexss}")
        print("*-" * 50)
        nums00, nums11 = 0, 0
        # gss = types.replace(" ", "").split(',')
        if types == "True":
            from win32com import client as wc
            # path_list = os.listdir(path)
            # doc_list = [os.path.join(path, str(i)) for i in path_list if str(i).endswith('doc')]
            doc_list = file_name(path, "doc")
            word = wc.Dispatch('Word.Application')
            print(f"需要转换的word有：{doc_list}")
            for path1 in doc_list:
                try:
                    save_path = str(path1).replace('doc', 'docx')
                    doc = word.Documents.Open(path1)
                    doc.SaveAs(save_path, 12, False, "", True, "", False, False, False, False)
                    doc.Close()
                    print(f'{save_path} 转存成功！')
                except Exception as e:
                    # print(f'{save_path} 转存失败，请手动转换')
                    print(e)
            word.Quit()
            print("*-" * 50)

        txtName = file_name(path, "docx")
        print("开始执行！")
        print("*-" * 50)
        for j in range(len(txtName)):
            print(f"处理文件{j}：{txtName[j]}")
            nums0, nums1 = write(txtName[j], tableIndexss, tableNums)
            nums00 += nums0
            nums11 += nums1
            print(f"文件{j}执行完毕，成功{nums0}，失败{nums1}")
            print("*" * 100)
            workbook.save(filename=save_file)
        print(f"文件共{len(txtName)}个，表格共成功{nums00}个，失败{nums11}个,请仔细校对")
    except Exception as e:
        print(e)
    os.system("pause")
